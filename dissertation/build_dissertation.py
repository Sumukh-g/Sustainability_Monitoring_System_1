"""Build the submission dissertation DOCX and PDF from chapter markdown."""
from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips, RGBColor
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parent
CH = ROOT / "chapters"
ASSETS = ROOT / "assets"
OUT_DOCX = ROOT / "AI_Based_Sustainability_Monitoring_System_Dissertation.docx"
OUT_PDF = ROOT / "AI_Based_Sustainability_Monitoring_System_Dissertation.pdf"

# Figure mapping: caption text -> image path
FIGURES = {
    "3.1": ASSETS / "diagrams" / "figure_3_1_system_context.png",
    "3.2": ASSETS / "diagrams" / "figure_3_2_component_architecture.png",
    "3.3": ASSETS / "diagrams" / "figure_3_3_ml_pipeline.png",
    "4.1": ASSETS / "figures" / "correlation_heatmap.png",
    "4.2": ASSETS / "figures" / "pue_trend.png",
    "4.3": ASSETS / "figures" / "wue_l_per_kwh_trend.png",
    "4.4": ASSETS / "figures" / "carbon_emissions_kg_trend.png",
    "4.5": ASSETS / "figures" / "total_energy_kwh_actual_vs_predicted.png",
    "4.6": ASSETS / "figures" / "total_energy_kwh_prediction_timeline.png",
    "4.7": ASSETS / "figures" / "total_energy_kwh_residual_distribution.png",
    "4.8": ASSETS / "figures" / "total_energy_kwh_residual_vs_predicted.png",
    "4.9": ASSETS / "screenshots" / "01_command_center.png",
    "4.10": ASSETS / "screenshots" / "02_energy_intelligence.png",
    "4.11": ASSETS / "screenshots" / "06_forecast_center.png",
    "4.12": ASSETS / "screenshots" / "07_anomaly_intelligence.png",
    "4.13": ASSETS / "screenshots" / "08_ai_advisor.png",
    "4.14": ASSETS / "screenshots" / "09_scenario_lab.png",
    "4.15": ASSETS / "screenshots" / "10_model_intelligence.png",
    "4.20": ASSETS / "figures" / "anomaly_confusion_matrix.png",
    "4.21": ASSETS / "figures" / "anomaly_timeline.png",
}


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_paragraph_format(p, space_after=8, space_before=0, line_spacing=1.5, align="justify"):
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_page_number(section, start_at=None, roman=False):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)
    set_run_font(run, size=11)
    sectPr = section._sectPr
    pgNumType = OxmlElement("w:pgNumType")
    if roman:
        pgNumType.set(qn("w:fmt"), "lowerRoman")
    if start_at is not None:
        pgNumType.set(qn("w:start"), str(start_at))
    # replace existing pgNumType if present
    for child in list(sectPr):
        if child.tag == qn("w:pgNumType"):
            sectPr.remove(child)
    sectPr.append(pgNumType)


def configure_section(section, different_first=False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.different_first_page_header_footer = different_first


def add_heading_custom(doc, text, level):
    sizes = {0: 20, 1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=14 if level else 0, space_after=10, align="left")
    if level == 0:
        p.paragraph_format.page_break_before = True
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 12), bold=True)
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text):
    # strip markdown bold/italic lightly
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph()
    set_paragraph_format(p)
    # handle **bold**
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = p.add_run(part[1:-1])
            set_run_font(run, italic=True)
        else:
            run = p.add_run(part)
            set_run_font(run)
    return p


def add_caption(doc, text, kind="Figure"):
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=4, space_after=10, align="center", line_spacing=1.15)
    run = p.add_run(text)
    set_run_font(run, size=11, italic=True)
    return p


def add_image(doc, path: Path, width=5.8):
    if not path.exists():
        add_body(doc, f"[Missing figure: {path.name}]")
        return
    p = doc.add_paragraph()
    set_paragraph_format(p, align="center", space_after=2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))


def parse_md_table(lines):
    rows = []
    for line in lines:
        if re.match(r"^\s*\|?\s*-+", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if cells:
            rows.append(cells)
    return rows


def add_table(doc, rows, caption=None):
    if caption:
        p = doc.add_paragraph()
        set_paragraph_format(p, space_before=8, space_after=4, align="left", line_spacing=1.15)
        run = p.add_run(caption)
        set_run_font(run, size=11, bold=True)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            cell_obj = table.cell(i, j)
            cell_obj.text = ""
            para = cell_obj.paragraphs[0]
            set_paragraph_format(para, space_after=2, space_before=2, line_spacing=1.0, align="left")
            run = para.add_run(cell)
            set_run_font(run, size=10, bold=(i == 0))
    doc.add_paragraph()


def insert_figure_if_marker(doc, line: str):
    # **Figure 3.1:** caption  OR  **Figure 4.1** ...
    m = re.match(r"^\*\*Figure\s+(\d+\.\d+):?\*\*\s*(.*)$", line.strip())
    if not m:
        m2 = re.match(r"^Figure\s+(\d+\.\d+):?\s*(.*)$", line.strip())
        if not m2:
            return False
        num, cap = m2.group(1), m2.group(2)
    else:
        num, cap = m.group(1), m.group(2)
    path = FIGURES.get(num)
    if path:
        add_image(doc, path, width=5.7 if "screenshot" in str(path) else 5.5)
    add_caption(doc, f"Figure {num}: {cap}".strip(": ").rstrip(".") + ".")
    return True


def render_markdown_file(doc, path: Path, chapter_level0_title: str | None = None):
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    # drop WORD_TARGET note and top H1 if we supply chapter title
    i = 0
    while i < len(lines) and (lines[i].startswith("WORD_TARGET") or not lines[i].strip()):
        i += 1
    if chapter_level0_title:
        add_heading_custom(doc, chapter_level0_title, 0)
        if i < len(lines) and lines[i].startswith("# "):
            i += 1
    table_buf = []
    in_table = False
    pending_table_caption = None

    while i < len(lines):
        line = lines[i]
        raw = line.rstrip()
        if not raw.strip():
            i += 1
            continue

        # table caption preceding table
        if re.match(r"^\*\*Table\s+\d+\.\d+:", raw) or re.match(r"^Table\s+\d+\.\d+:", raw):
            pending_table_caption = raw.replace("**", "")
            i += 1
            continue

        if "|" in raw and raw.strip().startswith("|"):
            in_table = True
            table_buf.append(raw)
            i += 1
            # continue collecting
            while i < len(lines) and "|" in lines[i]:
                table_buf.append(lines[i])
                i += 1
            rows = parse_md_table(table_buf)
            add_table(doc, rows, caption=pending_table_caption)
            pending_table_caption = None
            table_buf = []
            in_table = False
            continue

        if insert_figure_if_marker(doc, raw):
            i += 1
            continue

        if raw.startswith("### "):
            add_heading_custom(doc, raw[4:].strip(), 2)
        elif raw.startswith("## "):
            add_heading_custom(doc, raw[3:].strip(), 1)
        elif raw.startswith("# "):
            # already handled for chapters; treat as heading 0 if not
            if not chapter_level0_title:
                add_heading_custom(doc, raw[2:].strip(), 0)
        elif raw.startswith("- ") or raw.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_format(p, space_after=4, align="left", line_spacing=1.5)
            run = p.add_run(raw[2:].replace("**", ""))
            set_run_font(run)
        else:
            add_body(doc, raw)
        i += 1


def add_title_page(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_format(p, align="center", space_after=18)
    run = p.add_run(
        "An AI-Based Sustainability Monitoring System\nfor Data-Centre Operations"
    )
    set_run_font(run, size=20, bold=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, align="center", space_after=24)
    run = p.add_run(
        "A dissertation submitted in partial fulfilment of the requirements\n"
        "for the degree of [DEGREE TITLE]"
    )
    set_run_font(run, size=12)

    for label, value in [
        ("Student name", "[STUDENT NAME]"),
        ("Student ID", "[STUDENT ID]"),
        ("University", "[UNIVERSITY]"),
        ("School / Department", "[SCHOOL OR DEPARTMENT]"),
        ("Supervisor", "[SUPERVISOR NAME]"),
        ("Academic year", "[ACADEMIC YEAR]"),
        ("Submission date", "[SUBMISSION DATE]"),
    ]:
        p = doc.add_paragraph()
        set_paragraph_format(p, align="center", space_after=4, line_spacing=1.5)
        run = p.add_run(f"{label}: {value}")
        set_run_font(run, size=12)

    p = doc.add_paragraph()
    set_paragraph_format(p, align="center", space_before=36)
    run = p.add_run(
        "Repository analysed: https://github.com/Sumukh-g/Sustainability_Monitoring_System_1\n"
        "Commit: 756cc2493cdae4085620b385e7be4f3ce9cfc6af"
    )
    set_run_font(run, size=10, italic=True)


def add_toc_placeholder(doc):
    add_heading_custom(doc, "Table of Contents", 0)
    add_body(
        doc,
        "Update this table of contents in Microsoft Word (References → Table of Contents → Update Field) "
        "after opening the document so that page numbers refresh correctly.",
    )
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run2 = p.add_run("Right-click and select Update Field to generate the table of contents.")
    set_run_font(run2, italic=True, size=11)
    run._r.append(fldChar3)


def build():
    # ensure diagrams exist
    subprocess.run(
        [str(ROOT.parent / ".venv" / "Scripts" / "python.exe"), str(ROOT / "_make_diagrams.py")],
        check=False,
        cwd=str(ROOT.parent),
    )

    doc = Document()
    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.paragraph_format.line_spacing = 1.5

    section0 = doc.sections[0]
    configure_section(section0, different_first=True)
    add_page_number(section0, start_at=1, roman=True)

    add_title_page(doc)

    # Front matter
    render_markdown_file(doc, CH / "front_matter.md")
    add_heading_custom(doc, "Abstract", 0)
    # abstract body without H1
    abs_lines = (CH / "abstract.md").read_text(encoding="utf-8").splitlines()
    for line in abs_lines:
        if line.startswith("WORD_TARGET") or line.startswith("# ") or not line.strip():
            continue
        add_body(doc, line)

    # keywords if in front_matter already handled; ensure keywords paragraph
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=12)
    run = p.add_run(
        "Keywords: data-centre sustainability; PUE; forecasting; anomaly detection; decision support; Streamlit"
    )
    set_run_font(run, italic=True)

    render_markdown_file(doc, CH / "list_of_abbreviations.md")
    add_toc_placeholder(doc)

    # New section for Arabic numbering from Chapter 1
    new_section = doc.add_section()
    configure_section(new_section)
    add_page_number(new_section, start_at=1, roman=False)

    chapters = [
        (CH / "01_introduction.md", "CHAPTER 1: INTRODUCTION"),
        (CH / "02_literature.md", "CHAPTER 2: LITERATURE REVIEW AND RELATED WORK"),
        (CH / "03_methodology.md", "CHAPTER 3: METHODOLOGY, REQUIREMENTS AND SYSTEM DESIGN"),
        (CH / "04_results.md", "CHAPTER 4: IMPLEMENTATION, RESULTS, EVALUATION AND DISCUSSION"),
        (CH / "05_conclusion.md", "CHAPTER 5: CONCLUSION AND FUTURE WORK"),
    ]
    for path, title in chapters:
        render_markdown_file(doc, path, chapter_level0_title=title)

    # Extra screenshots appendix figures for Ch4 journey evidence
    add_heading_custom(doc, "Additional interface evidence", 1)
    add_body(
        doc,
        "The following screenshots were captured from the local Streamlit application during evaluation. "
        "They illustrate principal decision-support journeys used in Chapter 4. The runtime interface may "
        "display an internal product label; the dissertation refers to the artefact as the AI-Based "
        "Sustainability Monitoring System.",
    )
    for num, caption in [
        ("4.9", "Command Centre overview for London-DC1 and Manchester-DC2"),
        ("4.10", "Energy Intelligence page"),
        ("4.11", "Forecast Centre page"),
        ("4.12", "Anomaly Intelligence page"),
        ("4.13", "AI Advisor recommendations page"),
        ("4.14", "Scenario Lab what-if page"),
        ("4.15", "Model Intelligence comparison page"),
    ]:
        path = FIGURES[num]
        add_image(doc, path, width=5.6)
        add_caption(doc, f"Figure {num}: {caption}.")

    # Insert architecture figures if not already embedded via markers
    # Chapter 3 markers should insert 3.1; also add 3.2 and 3.3 after methodology if missing from md
    # They are referenced conceptually; ensure presence in appendix if needed.

    add_heading_custom(doc, "References", 0)
    ref_lines = (CH / "references.md").read_text(encoding="utf-8").splitlines()
    for line in ref_lines:
        if line.startswith("WORD_TARGET") or line.startswith("# ") or not line.strip():
            continue
        p = doc.add_paragraph()
        set_paragraph_format(p, space_after=6, align="left", line_spacing=1.15)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(line.strip())
        set_run_font(run, size=11)

    # Appendices
    add_heading_custom(doc, "APPENDICES", 0)
    add_heading_custom(doc, "Appendix A: Requirements and verification evidence", 1)
    add_body(
        doc,
        "The project maintains an executable requirements checklist and runtime audit. "
        "Mandatory requirements covering synthetic data generation, quality validation, KPI calculation, "
        "leakage-safe forecasting, Isolation Forest anomaly detection, recommendations, multi-page dashboard "
        "behaviour, automated tests and research artefacts were verified. Advanced items include the Scenario "
        "Lab simulator and renewable-energy context fields. Evidence files include reports/evaluation/*.csv, "
        "reports/figures/*.png, reports/FINAL_RESULTS.md, pytest output (12/12 passed) and validate_project.py "
        "(19/19 passed).",
    )
    add_heading_custom(doc, "Appendix B: Reproduction instructions", 1)
    for step in [
        "Create and activate a Python 3.11 virtual environment.",
        "Install dependencies from requirements.txt.",
        "Run python -m src.data_generation",
        "Run python -m src.forecasting",
        "Run python -m src.audit_pipeline",
        "Run pytest",
        "Run python validate_project.py",
        "Run streamlit run app.py",
    ]:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_format(p, space_after=4, align="left")
        run = p.add_run(step)
        set_run_font(run)

    add_heading_custom(doc, "Appendix C: Analysed repository commit", 1)
    add_body(
        doc,
        "Repository URL: https://github.com/Sumukh-g/Sustainability_Monitoring_System_1.git. "
        "Analysed commit hash: 756cc2493cdae4085620b385e7be4f3ce9cfc6af. "
        "Working-tree page filenames may differ from older documentation labels; evaluation used the runtime "
        "pages present in the local workspace under pages/.",
    )

    add_heading_custom(doc, "Appendix D: Supplementary architecture diagrams", 1)
    for num, cap in [
        ("3.2", "Component architecture of the monitoring artefact"),
        ("3.3", "Machine-learning training and inference data flow"),
    ]:
        add_image(doc, FIGURES[num], width=5.8)
        add_caption(doc, f"Figure {num}: {cap}.")

    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")

    # PDF conversion attempt
    try:
        from docx2pdf import convert

        convert(str(OUT_DOCX), str(OUT_PDF))
        print(f"Wrote {OUT_PDF}")
    except Exception as exc:
        print(f"docx2pdf failed: {exc}")
        # Word COM fallback
        try:
            import win32com.client  # type: ignore

            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            docw = word.Documents.Open(str(OUT_DOCX))
            docw.SaveAs(str(OUT_PDF), FileFormat=17)
            docw.Close()
            word.Quit()
            print(f"Wrote {OUT_PDF} via Word COM")
        except Exception as exc2:
            print(f"PDF conversion unavailable: {exc2}")


if __name__ == "__main__":
    build()
